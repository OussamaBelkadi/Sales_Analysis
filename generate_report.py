from docx import Document
from docx.shared import Inches
import os

# Créer un document Word
doc = Document()

# Titre principal
doc.add_heading("Rapport d’Analyse des Ventes", 0)

# Introduction
doc.add_paragraph(
    "Ce rapport présente une analyse complète des données de ventes extraites du fichier 'sales_data.xlsx'. "
    "Les indicateurs de performance, les ventes par produit, ville, client, canal et mois sont analysés."
)

# Ajouter un graphique
def add_section(title, image_path, doc):
    doc.add_heading(title, level=2)
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5.5))
    else:
        doc.add_paragraph("Image non trouvée : " + image_path)

# Ajouter tous les graphiques générés
add_section("1. Ventes par produit - comparaison annuelle", "output/ventes_par_produit_comparaison.png", doc)
add_section("2. Répartition des ventes - Top 5 Produits", "output/pie_top5_produits.png", doc)
add_section("3. Ventes par mois - comparaison annuelle", "output/ventes_par_mois_comparaison.png", doc)
add_section("4. Top 5 villes par ventes", "output/top_5_villes.png", doc)
add_section("5. Bénéfice par canal - comparaison annuelle", "output/profit_par_canal.png", doc)
add_section("6. Top 5 clients par ventes", "output/top_5_clients.png", doc)
add_section("7. Last 5 clients par ventes", "output/last_5_clients.png", doc)
add_section("8. KPI globaux", "output/kpi_cards.png", doc)

# Enregistrer le document
doc.save("output/rapport_analyse_ventes.pdf")

print("📄 Rapport Word généré avec succès dans output/rapport_analyse_ventes.docx ✅")
